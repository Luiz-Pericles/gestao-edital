from docx import Document
from io import BytesIO
import pandas as pd
from typing import List, Tuple

from schemas.excel_schema import validate_excel_columns, get_missing_columns
from utils.excel_classificacao import classificar_itens


class DocumentProcessor:
    @staticmethod
    def process_excel(excel_content: bytes) -> Tuple[pd.DataFrame, List[str]]:
        """
        Processa o arquivo Excel, valida as colunas e aplica a classificação.
        Retorna o DataFrame processado e lista de erros (se houver).
        """
        errors = []

        try:
            # Ler Excel
            df = pd.read_excel(BytesIO(excel_content))

            # Validar colunas
            if not validate_excel_columns(df.columns):
                missing = get_missing_columns(df.columns)
                errors.append(
                    f"Colunas obrigatórias faltando: {', '.join(missing)}")
                return None, errors

            # Aplicar classificação
            df_classificado = classificar_itens(df)
            return df_classificado, errors

        except Exception as e:
            errors.append(f"Erro ao processar planilha: {str(e)}")
            return None, errors

    @staticmethod
    def create_word_table(doc: Document, df: pd.DataFrame) -> Document:
        """
        Cria uma nova tabela Word a partir do DataFrame.
        """
        # Criar tabela com cabeçalho
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Adicionar cabeçalhos
        for j, column in enumerate(df.columns):
            table.cell(0, j).text = str(column)

        # Adicionar dados
        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value)

        return doc

    @staticmethod
    def process_documents(word_content: bytes, excel_content: bytes, selected_indices: List[int]) -> Tuple[BytesIO, List[str]]:
        """
        Processa os documentos Word e Excel, substituindo as tabelas selecionadas.
        Retorna o documento modificado e lista de erros (se houver).
        """
        errors = []

        try:
            # Processar Excel
            df_classificado, excel_errors = DocumentProcessor.process_excel(
                excel_content)
            if excel_errors:
                return None, excel_errors

            # Processar Word
            doc = Document(BytesIO(word_content))

            # Remover tabelas selecionadas
            removed = False
            for i, table in enumerate(doc.tables[:]):
                if i in selected_indices:
                    parent = table._element.getparent()
                    if parent is not None:
                        parent.remove(table._element)
                        if not removed:  # Inserir nova tabela após remover a primeira selecionada
                            doc = DocumentProcessor.create_word_table(
                                doc, df_classificado)
                            removed = True

            # Salvar documento
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            return output, errors

        except Exception as e:
            errors.append(f"Erro ao processar documentos: {str(e)}")
            return None, errors
